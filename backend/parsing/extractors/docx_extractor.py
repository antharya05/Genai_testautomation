"""DOCX extractor — recovers BOTH paragraphs and tables.

The previous implementation in ``main.py`` read paragraphs only and silently
dropped every table, which is exactly where automotive requirements usually
live. Here tables are preserved as normalized ``list[list[str]]`` for the
structured stage, and paragraph text is kept for the semi-structured fallback.
"""

from __future__ import annotations

from . import ExtractedContent, Table


def extract(path: str) -> ExtractedContent:
    from docx import Document

    doc = Document(path)
    content = ExtractedContent()

    # Paragraphs (for semi-structured fallback).
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    content.text = "\n".join(paras)

    # Tables (for structured parsing).
    for table in doc.tables:
        rows: Table = []
        for row in table.rows:
            cells = [_cell_text(cell) for cell in row.cells]
            rows.append(cells)
        rows = _dedupe_merged_cells(rows)
        if rows:
            content.tables.append(rows)

    return content


def _cell_text(cell) -> str:
    # A cell may contain multiple paragraphs; join them with spaces.
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()).strip()


def _dedupe_merged_cells(rows: Table) -> Table:
    """python-docx repeats the same cell object across a horizontal merge, so a
    merged header can appear as ['ID','ID','ID']. Collapse consecutive
    duplicates only when an entire row is a single repeated value."""
    cleaned: Table = []
    for row in rows:
        if row and len(set(row)) == 1 and len(row) > 1:
            cleaned.append([row[0]])
        else:
            cleaned.append(row)
    return cleaned
